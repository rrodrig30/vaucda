"""
Document Parsing Utilities
Extracts text from PDF, DOCX, and TXT files for RAG pipeline
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional
import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Parse various document formats for RAG ingestion.

    Supports:
    - PDF files (.pdf)
    - Microsoft Word documents (.docx)
    - Plain text files (.txt)
    """

    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse document and extract text with metadata.

        Args:
            file_path: Path to document file

        Returns:
            Dictionary containing:
            - text: Extracted text content
            - metadata: File metadata (name, type, pages, etc.)

        Raises:
            ValueError: If file type is unsupported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = path.suffix.lower()

        if file_ext == '.pdf':
            return self._parse_pdf(path)
        elif file_ext == '.docx':
            return self._parse_docx(path)
        elif file_ext == '.txt':
            return self._parse_txt(path)
        else:
            raise ValueError(
                f"Unsupported file type: {file_ext}. "
                "Supported: .pdf, .docx, .txt"
            )

    def _parse_pdf(self, path: Path) -> Dict[str, Any]:
        """Extract text from PDF file."""
        try:
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                # Extract text from all pages
                text_parts = []
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text_parts.append(page.extract_text())

                text = "\n\n".join(text_parts)

                # Extract metadata
                metadata = {
                    'filename': path.name,
                    'file_type': 'pdf',
                    'num_pages': num_pages,
                    'file_size': path.stat().st_size,
                }

                # Add PDF metadata if available
                if pdf_reader.metadata:
                    if pdf_reader.metadata.title:
                        metadata['title'] = pdf_reader.metadata.title
                    if pdf_reader.metadata.author:
                        metadata['author'] = pdf_reader.metadata.author
                    if pdf_reader.metadata.subject:
                        metadata['subject'] = pdf_reader.metadata.subject

                logger.info(
                    f"Parsed PDF: {path.name} ({num_pages} pages, "
                    f"{len(text)} chars)"
                )

                return {
                    'text': text,
                    'metadata': metadata
                }

        except Exception as e:
            logger.error(f"Failed to parse PDF {path.name}: {e}")
            raise

    def _parse_docx(self, path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(path)

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            text = "\n\n".join(text_parts)

            # Build metadata
            metadata = {
                'filename': path.name,
                'file_type': 'docx',
                'num_paragraphs': len(doc.paragraphs),
                'num_tables': len(doc.tables),
                'file_size': path.stat().st_size,
            }

            # Add core properties if available
            core_props = doc.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.created:
                metadata['created_date'] = core_props.created.isoformat()

            logger.info(
                f"Parsed DOCX: {path.name} ({len(doc.paragraphs)} paragraphs, "
                f"{len(text)} chars)"
            )

            return {
                'text': text,
                'metadata': metadata
            }

        except Exception as e:
            logger.error(f"Failed to parse DOCX {path.name}: {e}")
            raise

    def _parse_txt(self, path: Path) -> Dict[str, Any]:
        """Extract text from TXT file."""
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                with open(path, 'r', encoding='utf-8') as file:
                    text = file.read()
            except UnicodeDecodeError:
                logger.warning(f"UTF-8 decode failed for {path.name}, trying latin-1")
                with open(path, 'r', encoding='latin-1') as file:
                    text = file.read()

            # Count lines
            num_lines = text.count('\n') + 1

            metadata = {
                'filename': path.name,
                'file_type': 'txt',
                'num_lines': num_lines,
                'file_size': path.stat().st_size,
            }

            logger.info(
                f"Parsed TXT: {path.name} ({num_lines} lines, {len(text)} chars)"
            )

            return {
                'text': text,
                'metadata': metadata
            }

        except Exception as e:
            logger.error(f"Failed to parse TXT {path.name}: {e}")
            raise


# Global singleton
_document_parser = None


def get_document_parser() -> DocumentParser:
    """Get or create global document parser instance."""
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
    return _document_parser
