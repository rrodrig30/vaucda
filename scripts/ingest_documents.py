#!/usr/bin/env python3
"""
Document Ingestion Pipeline for VAUCDA Knowledge Base

Ingests medical documents (PDF, DOCX) into Neo4j with embeddings:
1. Parse documents (PDF/DOCX)
2. Chunk using appropriate strategy
3. Generate embeddings
4. Store in Neo4j with graph relationships

Usage:
    python scripts/ingest_documents.py --directory /path/to/docs --doc-type guideline
    python scripts/ingest_documents.py --file /path/to/document.pdf --doc-type literature
"""

import asyncio
import argparse
import logging
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional
import json

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from backend.rag.chunking import MedicalDocumentChunker, DocumentType, DocumentChunk
from backend.rag.embeddings import EmbeddingGenerator
from backend.database.neo4j_client import Neo4jClient, Neo4jConfig

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse PDF and DOCX documents."""

    @staticmethod
    def parse_pdf(file_path: Path) -> str:
        """Parse PDF document to text."""
        try:
            import PyPDF2

            text_parts = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)

                logger.info(f"Parsing PDF: {file_path.name} ({num_pages} pages)")

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    text_parts.append(text)

                    if page_num % 10 == 0:
                        logger.info(f"  Processed {page_num}/{num_pages} pages")

            full_text = "\n\n".join(text_parts)
            logger.info(f"PDF parsed: {len(full_text)} characters")

            return full_text

        except ImportError:
            logger.error("PyPDF2 not installed. Install: pip install PyPDF2")
            raise
        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
            raise

    @staticmethod
    def parse_docx(file_path: Path) -> str:
        """Parse DOCX document to text."""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            logger.info(f"Parsing DOCX: {file_path.name}")

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            full_text = "\n\n".join(text_parts)
            logger.info(f"DOCX parsed: {len(full_text)} characters")

            return full_text

        except ImportError:
            logger.error("python-docx not installed. Install: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise

    @staticmethod
    def parse_txt(file_path: Path) -> str:
        """Parse plain text file."""
        try:
            logger.info(f"Parsing TXT: {file_path.name}")
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            logger.info(f"TXT parsed: {len(text)} characters")
            return text
        except Exception as e:
            logger.error(f"Failed to parse TXT {file_path}: {e}")
            raise


class DocumentIngestionPipeline:
    """Complete document ingestion pipeline."""

    def __init__(
        self,
        neo4j_client: Neo4jClient,
        embedding_generator: EmbeddingGenerator,
        chunker: MedicalDocumentChunker
    ):
        """
        Initialize ingestion pipeline.

        Args:
            neo4j_client: Neo4j client for storage
            embedding_generator: Embedding generator
            chunker: Document chunker
        """
        self.neo4j = neo4j_client
        self.embedder = embedding_generator
        self.chunker = chunker
        self.parser = DocumentParser()

    async def ingest_file(
        self,
        file_path: Path,
        doc_type: DocumentType,
        metadata: Optional[Dict[str, Any]] = None
    ) -> int:
        """
        Ingest a single document file.

        Args:
            file_path: Path to document
            doc_type: Type of document
            metadata: Additional metadata

        Returns:
            Number of chunks ingested
        """
        logger.info(f"Ingesting file: {file_path}")

        # 1. Parse document
        if file_path.suffix.lower() == '.pdf':
            text = self.parser.parse_pdf(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            text = self.parser.parse_docx(file_path)
        elif file_path.suffix.lower() == '.txt':
            text = self.parser.parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

        # 2. Build metadata
        doc_metadata = {
            "source_file": file_path.name,
            "title": file_path.stem.replace('_', ' ').title(),
            "document_type": doc_type.value,
            **(metadata or {})
        }

        # 3. Chunk document
        logger.info(f"Chunking document (type: {doc_type.value})...")
        chunks = self.chunker.chunk_document(text, doc_type, doc_metadata)
        logger.info(f"Created {len(chunks)} chunks")

        # 4. Generate embeddings
        logger.info("Generating embeddings...")
        chunk_texts = [chunk.content for chunk in chunks]
        embeddings = self.embedder.generate_embeddings_batch(
            chunk_texts,
            batch_size=32,
            show_progress=True
        )

        # Attach embeddings to chunks
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding.tolist()

        # 5. Prepare for Neo4j ingestion
        logger.info("Preparing documents for storage...")
        documents = []
        for chunk in chunks:
            doc = {
                "id": f"{doc_metadata['source_file']}_{chunk.chunk_index}",
                "title": doc_metadata["title"],
                "content": chunk.content,
                "summary": chunk.content[:200] + "..." if len(chunk.content) > 200 else chunk.content,
                "embedding": chunk.embedding,
                "category": doc_metadata.get("category", "general"),
                "source": doc_metadata.get("source", "Unknown"),
                "document_type": doc_type.value,
                "version": doc_metadata.get("version", "1.0"),
                "publication_date": doc_metadata.get("publication_date", "2024-01-01"),
                "keywords": doc_metadata.get("keywords", []),
                "chunk_index": chunk.chunk_index,
                "total_chunks": chunk.total_chunks
            }
            documents.append(doc)

        # 6. Ingest into Neo4j
        logger.info("Ingesting into Neo4j...")
        ingested_count = await self.neo4j.batch_ingest_documents(
            documents,
            batch_size=100
        )

        logger.info(f"Successfully ingested {ingested_count} chunks from {file_path.name}")
        return ingested_count

    async def ingest_directory(
        self,
        directory: Path,
        doc_type: DocumentType,
        metadata: Optional[Dict[str, Any]] = None,
        recursive: bool = False
    ) -> int:
        """
        Ingest all documents in a directory.

        Args:
            directory: Directory containing documents
            doc_type: Type of documents
            metadata: Common metadata for all documents
            recursive: Whether to search recursively

        Returns:
            Total number of chunks ingested
        """
        logger.info(f"Ingesting directory: {directory}")

        # Find all supported files
        patterns = ['*.pdf', '*.docx', '*.doc', '*.txt']
        files = []

        for pattern in patterns:
            if recursive:
                files.extend(directory.rglob(pattern))
            else:
                files.extend(directory.glob(pattern))

        logger.info(f"Found {len(files)} documents to ingest")

        total_chunks = 0
        for idx, file_path in enumerate(files, 1):
            try:
                logger.info(f"\n[{idx}/{len(files)}] Processing: {file_path.name}")
                chunks = await self.ingest_file(file_path, doc_type, metadata)
                total_chunks += chunks
            except Exception as e:
                logger.error(f"Failed to ingest {file_path}: {e}")
                continue

        logger.info(f"\nDirectory ingestion complete: {total_chunks} total chunks from {len(files)} files")
        return total_chunks


async def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description="Ingest medical documents into VAUCDA knowledge base"
    )

    parser.add_argument(
        '--file',
        type=str,
        help='Path to single document file'
    )
    parser.add_argument(
        '--directory',
        type=str,
        help='Path to directory containing documents'
    )
    parser.add_argument(
        '--doc-type',
        type=str,
        required=True,
        choices=['guideline', 'calculator', 'literature', 'general'],
        help='Type of documents'
    )
    parser.add_argument(
        '--category',
        type=str,
        help='Document category (prostate, kidney, etc.)'
    )
    parser.add_argument(
        '--source',
        type=str,
        help='Source organization (AUA, NCCN, etc.)'
    )
    parser.add_argument(
        '--version',
        type=str,
        default='1.0',
        help='Document version'
    )
    parser.add_argument(
        '--publication-date',
        type=str,
        help='Publication date (YYYY-MM-DD)'
    )
    parser.add_argument(
        '--recursive',
        action='store_true',
        help='Search directory recursively'
    )
    parser.add_argument(
        '--metadata-file',
        type=str,
        help='JSON file containing additional metadata'
    )

    args = parser.parse_args()

    # Validate arguments
    if not args.file and not args.directory:
        parser.error("Must specify either --file or --directory")

    # Load metadata from file if provided
    extra_metadata = {}
    if args.metadata_file:
        with open(args.metadata_file, 'r') as f:
            extra_metadata = json.load(f)

    # Build metadata
    metadata = {
        "category": args.category or "general",
        "source": args.source or "Unknown",
        "version": args.version,
        "publication_date": args.publication_date or "2024-01-01",
        **extra_metadata
    }

    # Initialize components
    logger.info("Initializing ingestion pipeline...")

    try:
        neo4j_config = Neo4jConfig()
        neo4j_client = Neo4jClient(neo4j_config)

        # Verify Neo4j connectivity
        if not await neo4j_client.verify_connectivity():
            logger.error("Failed to connect to Neo4j. Ensure Neo4j is running.")
            return 1

        embedding_generator = EmbeddingGenerator()
        chunker = MedicalDocumentChunker()

        pipeline = DocumentIngestionPipeline(
            neo4j_client=neo4j_client,
            embedding_generator=embedding_generator,
            chunker=chunker
        )

        # Convert doc_type string to enum
        doc_type = DocumentType(args.doc_type)

        # Ingest
        if args.file:
            file_path = Path(args.file)
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return 1

            await pipeline.ingest_file(file_path, doc_type, metadata)

        elif args.directory:
            dir_path = Path(args.directory)
            if not dir_path.exists():
                logger.error(f"Directory not found: {dir_path}")
                return 1

            await pipeline.ingest_directory(
                dir_path,
                doc_type,
                metadata,
                recursive=args.recursive
            )

        logger.info("Ingestion complete!")

        # Close connections
        await neo4j_client.close()

        return 0

    except Exception as e:
        logger.error(f"Ingestion failed: {e}", exc_info=True)
        return 1


if __name__ == "__main__":
    exit_code = asyncio.run(main())
    sys.exit(exit_code)
